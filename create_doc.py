from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
doc.add_heading('Galaxy ISO Audit Management System - User Guide', 0)
doc.add_paragraph('Welcome to the Galaxy ISO Audit Management System. This guide explains all features and how to use them effectively.')

# System Overview
doc.add_heading('System Overview', 1)
doc.add_paragraph('This system helps organizations manage audits following ISO 19011 guidelines. It covers the complete audit lifecycle from planning to follow-up, with built-in compliance tracking, risk management, and corrective action management.')

# User Roles
doc.add_heading('User Roles', 1)
doc.add_paragraph('The system has 6 user roles with different access levels:')

table = doc.add_table(rows=7, cols=3)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = 'Role'
hdr[1].text = 'Description'
hdr[2].text = 'Access Level'

rows_data = [
    ['System Admin', 'Full system access. Manages users, departments, and system settings', 'All features'],
    ['Audit Manager', 'Plans and oversees audits. Assigns auditors and reviews reports', 'Planning, Analytics, Assets, Vendors, Access Control'],
    ['Auditor', 'Conducts audits, collects evidence, documents findings', 'Audits, Evidence, Findings, Reports'],
    ['Department Head', 'Reviews audit findings for their department. Approves corrective actions', 'View audits, Approve workflows, CAPA'],
    ['Department Officer', 'Responds to audit queries. Provides evidence and implements actions', 'View assigned audits, Upload evidence, Respond to queries'],
    ['Viewer', 'Read-only access to audit information', 'View only']
]
for i, row_data in enumerate(rows_data):
    row = table.rows[i+1].cells
    row[0].text = row_data[0]
    row[1].text = row_data[1]
    row[2].text = row_data[2]

doc.add_paragraph()

# Sidebar Features
doc.add_heading('Sidebar Features', 1)

# 1. Dashboard
doc.add_heading('1. Dashboard', 2)
doc.add_paragraph('What it does: Shows a summary of your audit program at a glance.')
doc.add_paragraph('You will see:')
doc.add_paragraph('- Total audits and their status (planned, executing, closed)')
doc.add_paragraph('- Open findings by severity (critical, high, medium, low)')
doc.add_paragraph('- Risk heatmap showing risk distribution')
doc.add_paragraph('- Compliance scores across ISO frameworks')
doc.add_paragraph('- CAPA (Corrective Action) status tracker')
doc.add_paragraph('- Overdue follow-ups')
doc.add_paragraph('ISO Reference: ISO 19011 Clause 5.5 - Monitoring the audit programme')

# 2. Audits
doc.add_heading('2. Audits', 2)
doc.add_paragraph('What it does: Create and manage individual audits through their complete lifecycle.')
doc.add_paragraph('Audit Lifecycle (ISO 19011 Clause 6):')

table2 = doc.add_table(rows=8, cols=3)
table2.style = 'Table Grid'
hdr2 = table2.rows[0].cells
hdr2[0].text = 'Phase'
hdr2[1].text = 'ISO Clause'
hdr2[2].text = 'What Happens'
audit_rows = [
    ['Planned', '6.2', 'Audit is scheduled but not started'],
    ['Initiated', '6.2', 'Define objectives, scope, criteria, and assign team'],
    ['Preparation', '6.3', 'Create checklists, request documents, plan interviews'],
    ['Executing', '6.4', 'Collect evidence, conduct interviews, document observations'],
    ['Reporting', '6.5', 'Generate audit report with findings and recommendations'],
    ['Follow-up', '6.6', 'Track corrective actions and verify implementation'],
    ['Closed', '-', 'Audit complete, all actions verified']
]
for i, row_data in enumerate(audit_rows):
    row = table2.rows[i+1].cells
    row[0].text = row_data[0]
    row[1].text = row_data[1]
    row[2].text = row_data[2]

doc.add_paragraph()
doc.add_paragraph('How to use:')
doc.add_paragraph('1. Click "Create Audit" to start a new audit')
doc.add_paragraph('2. Fill in title, year, scope, and risk rating')
doc.add_paragraph('3. Follow the guided workflow through each phase')
doc.add_paragraph('4. Use the navigation tabs to access different audit sections')

doc.add_paragraph('Audit Sub-sections:')
doc.add_paragraph('- Initiate - Set objectives, scope, criteria, methodology')
doc.add_paragraph('- Team - Assign lead auditor and team members')
doc.add_paragraph('- Prepare - Create checklists and document requests')
doc.add_paragraph('- Work Program - Define audit procedures')
doc.add_paragraph('- Evidence - Upload and manage audit evidence')
doc.add_paragraph('- Findings - Document non-conformities and observations')
doc.add_paragraph('- Queries - Communication between auditors and auditees')
doc.add_paragraph('- Report - Generate and approve audit reports')
doc.add_paragraph('- Follow-up - Track corrective action implementation')

# 3. Workflows
doc.add_heading('3. Workflows', 2)
doc.add_paragraph('What it does: Manages approval processes for audit reports and documents.')
doc.add_paragraph('How it works:')
doc.add_paragraph('1. Create a workflow linked to an audit')
doc.add_paragraph('2. Add approval steps with assigned departments/users')
doc.add_paragraph('3. Each step must be approved before moving to the next')
doc.add_paragraph('4. Track workflow status: Pending → In Progress → Approved/Rejected')
doc.add_paragraph('Actions available:')
doc.add_paragraph('- Approve - Accept and move to next step')
doc.add_paragraph('- Reject - Send back with comments')
doc.add_paragraph('- Return - Request changes')
doc.add_paragraph('- Sign - Add digital signature')
doc.add_paragraph('- Review - Mark as reviewed')
doc.add_paragraph('- Acknowledge - Confirm receipt')
doc.add_paragraph('ISO Reference: ISO 19011 Clause 6.5.2 - Approving and distributing the audit report')

# 4. Planning
doc.add_heading('4. Planning', 2)
doc.add_paragraph('What it does: Plan your annual audit schedule.')
doc.add_paragraph('Features:')
doc.add_paragraph('- Create annual audit programmes')
doc.add_paragraph('- Schedule audits by quarter')
doc.add_paragraph('- Risk-based audit selection')
doc.add_paragraph('- Resource allocation')
doc.add_paragraph('ISO Reference: ISO 19011 Clause 5 - Managing an audit programme')
doc.add_paragraph('Access: System Admin, Audit Manager only')

# 5. Reports
doc.add_heading('5. Reports', 2)
doc.add_paragraph('What it does: Generate and manage audit reports.')
doc.add_paragraph('Report types:')
doc.add_paragraph('- Individual audit reports')
doc.add_paragraph('- Summary reports')
doc.add_paragraph('- Compliance reports')
doc.add_paragraph('- Trend analysis reports')
doc.add_paragraph('Features:')
doc.add_paragraph('- AI-powered report generation')
doc.add_paragraph('- Version control')
doc.add_paragraph('- Review and approval workflow')
doc.add_paragraph('- Export to PDF/Word')
doc.add_paragraph('ISO Reference: ISO 19011 Clause 6.5 - Preparing and distributing the audit report')

# 6. Follow-ups
doc.add_heading('6. Follow-ups', 2)
doc.add_paragraph('What it does: Track corrective actions after audits.')
doc.add_paragraph('How to use:')
doc.add_paragraph('1. View all follow-up items across audits')
doc.add_paragraph('2. Filter by status: Pending, In Progress, Completed, Overdue')
doc.add_paragraph('3. Assign responsible persons')
doc.add_paragraph('4. Set due dates')
doc.add_paragraph('5. Upload evidence of completion')
doc.add_paragraph('6. Verify implementation')
doc.add_paragraph('ISO Reference: ISO 19011 Clause 6.6 - Conducting audit follow-up')

# 7. Risk Assessment
doc.add_heading('7. Risk Assessment', 2)
doc.add_paragraph('What it does: Identify, assess, and manage risks.')
doc.add_paragraph('Features:')
doc.add_paragraph('- Create risk assessments linked to audits or assets')
doc.add_paragraph('- Risk matrix visualization (5x5 grid)')
doc.add_paragraph('- Calculate risk scores (Likelihood × Impact)')
doc.add_paragraph('- Define mitigation plans')
doc.add_paragraph('- Link controls to risks')
doc.add_paragraph('- AI-suggested controls based on ISO 27001')
doc.add_paragraph('Risk Categories:')

table3 = doc.add_table(rows=5, cols=2)
table3.style = 'Table Grid'
hdr3 = table3.rows[0].cells
hdr3[0].text = 'Score'
hdr3[1].text = 'Category'
risk_rows = [['1-4', 'Low'], ['5-9', 'Medium'], ['10-15', 'High'], ['16-25', 'Critical']]
for i, row_data in enumerate(risk_rows):
    row = table3.rows[i+1].cells
    row[0].text = row_data[0]
    row[1].text = row_data[1]

doc.add_paragraph()
doc.add_paragraph('ISO Reference: ISO 31000 - Risk Management')

# 8. CAPA Management
doc.add_heading('8. CAPA Management', 2)
doc.add_paragraph('What it does: Manage Corrective and Preventive Actions.')
doc.add_paragraph('CAPA Types:')
doc.add_paragraph('- Corrective - Fix existing problems')
doc.add_paragraph('- Preventive - Prevent future problems')
doc.add_paragraph('- Both - Combined approach')
doc.add_paragraph('CAPA Workflow:')
doc.add_paragraph('1. Open - CAPA created')
doc.add_paragraph('2. In Progress - Actions being implemented')
doc.add_paragraph('3. Pending Verification - Awaiting effectiveness check')
doc.add_paragraph('4. Closed - Verified and complete')
doc.add_paragraph('Features:')
doc.add_paragraph('- Root cause analysis (5 Whys, Fishbone, etc.)')
doc.add_paragraph('- Link to findings and risks')
doc.add_paragraph('- Track progress percentage')
doc.add_paragraph('- Effectiveness review')
doc.add_paragraph('- Cost tracking')
doc.add_paragraph('ISO Reference: ISO 9001 Clause 10.2 - Nonconformity and corrective action')

# 9. Documents
doc.add_heading('9. Documents', 2)
doc.add_paragraph('What it does: Central document library with version control.')
doc.add_paragraph('Document lifecycle:')
doc.add_paragraph('- Draft → Under Review → Approved → Active → Expired/Archived')
doc.add_paragraph('Features:')
doc.add_paragraph('- Upload documents with metadata')
doc.add_paragraph('- Version control')
doc.add_paragraph('- Approval workflow')
doc.add_paragraph('- Expiry tracking')
doc.add_paragraph('- Confidentiality levels (Public, Internal, Confidential, Restricted)')
doc.add_paragraph('- Search and filter')
doc.add_paragraph('ISO Reference: ISO 9001 Clause 7.5 - Documented information')

# 10. Assets
doc.add_heading('10. Assets', 2)
doc.add_paragraph('What it does: Manage organizational assets for audit scope.')
doc.add_paragraph('Asset categories:')
doc.add_paragraph('- Hardware')
doc.add_paragraph('- Software')
doc.add_paragraph('- Data')
doc.add_paragraph('- People')
doc.add_paragraph('- Facilities')
doc.add_paragraph('- Services')
doc.add_paragraph('Features:')
doc.add_paragraph('- Asset inventory')
doc.add_paragraph('- Criticality assessment')
doc.add_paragraph('- Link assets to risks')
doc.add_paragraph('- Track asset owners')
doc.add_paragraph('ISO Reference: ISO 27001 Annex A.8 - Asset management')
doc.add_paragraph('Access: System Admin, Audit Manager only')

# 11. Vendors
doc.add_heading('11. Vendors', 2)
doc.add_paragraph('What it does: Manage third-party vendors and suppliers.')
doc.add_paragraph('Features:')
doc.add_paragraph('- Vendor registry')
doc.add_paragraph('- Risk rating (Low, Medium, High, Critical)')
doc.add_paragraph('- Compliance tracking')
doc.add_paragraph('- Contract management')
doc.add_paragraph('- Performance monitoring')
doc.add_paragraph('ISO Reference: ISO 27001 Annex A.15 - Supplier relationships')
doc.add_paragraph('Access: System Admin, Audit Manager only')

# 12. Analytics
doc.add_heading('12. Analytics', 2)
doc.add_paragraph('What it does: Advanced reporting and trend analysis.')
doc.add_paragraph('Reports available:')
doc.add_paragraph('- Audit completion trends')
doc.add_paragraph('- Finding trends by severity')
doc.add_paragraph('- Compliance score trends')
doc.add_paragraph('- Risk distribution analysis')
doc.add_paragraph('- CAPA effectiveness metrics')
doc.add_paragraph('- Department performance')
doc.add_paragraph('ISO Reference: ISO 19011 Clause 5.6 - Reviewing and improving the audit programme')
doc.add_paragraph('Access: System Admin, Audit Manager only')

# 13. Users
doc.add_heading('13. Users', 2)
doc.add_paragraph('What it does: Manage system users.')
doc.add_paragraph('Features:')
doc.add_paragraph('- Create/edit users')
doc.add_paragraph('- Assign roles')
doc.add_paragraph('- Assign to departments')
doc.add_paragraph('- Activate/deactivate accounts')
doc.add_paragraph('Access: System Admin only')

# 14. Departments
doc.add_heading('14. Departments', 2)
doc.add_paragraph('What it does: Manage organizational structure.')
doc.add_paragraph('Features:')
doc.add_paragraph('- Create departments')
doc.add_paragraph('- Set parent-child relationships')
doc.add_paragraph('- Assign department heads')
doc.add_paragraph('Access: System Admin only')

# 15. Access Control
doc.add_heading('15. Access Control', 2)
doc.add_paragraph('What it does: Fine-grained permission management with enhanced Role-Based Access Control (RBAC).')
doc.add_paragraph('The Access Control module has 4 sub-features (tabs):')
doc.add_paragraph('Access: System Admin, Audit Manager only')
doc.add_paragraph('ISO References:')
doc.add_paragraph('- ISO 27001 A.9.2.2 - User access provisioning')
doc.add_paragraph('- ISO 27001 A.6.1.2 - Segregation of duties')
doc.add_paragraph('- ISO 27001 A.12.4.1 - Event logging')

# Sub-Feature 1: Team Assignment
doc.add_heading('Sub-Feature 1: Team Assignment', 3)
doc.add_paragraph('What it does: Assign auditors to audit teams with proper roles per ISO 19011 requirements.')
doc.add_paragraph('How to use:')
doc.add_paragraph('1. Enter or select an Audit ID')
doc.add_paragraph('2. Select a Lead Auditor from available users')
doc.add_paragraph('3. Add team members and assign their role in the audit')
doc.add_paragraph('4. Click "Assign Team"')
doc.add_paragraph('Audit Team Roles:')

table4 = doc.add_table(rows=7, cols=2)
table4.style = 'Table Grid'
hdr4 = table4.rows[0].cells
hdr4[0].text = 'Role'
hdr4[1].text = 'Description'
team_rows = [
    ['Lead Auditor', 'Heads the audit team, responsible for overall audit'],
    ['Senior Auditor', 'Experienced auditor, can mentor others'],
    ['Auditor', 'Conducts audit procedures and collects evidence'],
    ['Technical Specialist', 'Provides expertise in specific technical areas'],
    ['Observer', 'Watches the audit process (for training/oversight)'],
    ['Trainee Auditor', 'Learning auditor under supervision']
]
for i, row_data in enumerate(team_rows):
    row = table4.rows[i+1].cells
    row[0].text = row_data[0]
    row[1].text = row_data[1]

doc.add_paragraph()
doc.add_paragraph('Team Statistics shown:')
doc.add_paragraph('- Total team size')
doc.add_paragraph('- Number of auditors by role type')
doc.add_paragraph('- Competency verification status')
doc.add_paragraph('ISO 19011 Compliance:')
doc.add_paragraph('- Team competency must be verified before audit execution')
doc.add_paragraph('- Lead auditor must have appropriate qualifications')
doc.add_paragraph('- Team composition should match audit scope and complexity')

# Sub-Feature 2: User Management
doc.add_heading('Sub-Feature 2: User Management', 3)
doc.add_paragraph('What it does: Manage users and assign additional roles from the role matrix.')
doc.add_paragraph('How to use:')
doc.add_paragraph('1. Search or filter users by name, role, or department')
doc.add_paragraph('2. Click on a user to select them')
doc.add_paragraph('3. View their current details and role assignments')
doc.add_paragraph('4. Assign additional roles if needed')
doc.add_paragraph('User Details shown:')
doc.add_paragraph('- Name and email')
doc.add_paragraph('- Primary role (system role)')
doc.add_paragraph('- Department')
doc.add_paragraph('- Active/Inactive status')
doc.add_paragraph('Assigning Additional Roles:')
doc.add_paragraph('1. Select a role from the Role Matrix dropdown')
doc.add_paragraph('2. Enter a reason for the assignment')
doc.add_paragraph('3. Check "Temporary Assignment" if time-limited')
doc.add_paragraph('4. Set expiry date for temporary roles')
doc.add_paragraph('5. Click "Assign Role"')
doc.add_paragraph('Role Assignment Features:')
doc.add_paragraph('- Multiple roles can be assigned to one user')
doc.add_paragraph('- Temporary assignments auto-expire')
doc.add_paragraph('- Assignment reasons are logged for audit trail')
doc.add_paragraph('- View all current and historical assignments')
doc.add_paragraph('Why use additional roles?')
doc.add_paragraph('- Give a user extra permissions for a specific project')
doc.add_paragraph('- Temporary access during staff absence')
doc.add_paragraph('- Cross-department collaboration needs')

# Sub-Feature 3: Audit Visibility
doc.add_heading('Sub-Feature 3: Audit Visibility', 3)
doc.add_paragraph('What it does: Controls which audits users can see based on their role and department.')
doc.add_paragraph('Your Audit Access section shows:')
doc.add_paragraph('- Number of audits you can access')
doc.add_paragraph('- Your access level (Full, Department, Assigned Only, None)')
doc.add_paragraph('- List of accessible audits with status and access reason')
doc.add_paragraph('Access Levels Explained:')

table5 = doc.add_table(rows=5, cols=3)
table5.style = 'Table Grid'
hdr5 = table5.rows[0].cells
hdr5[0].text = 'Level'
hdr5[1].text = 'Who Gets It'
hdr5[2].text = 'What They See'
access_rows = [
    ['Full Access', 'System Administrators', 'All audits in the system'],
    ['Department + Assigned', 'Audit Managers', 'Audits in their department + audits assigned to them'],
    ['Assigned Only', 'Auditors', 'Only audits they are assigned to'],
    ['Department Audits', 'Department Staff', 'Audits related to their department']
]
for i, row_data in enumerate(access_rows):
    row = table5.rows[i+1].cells
    row[0].text = row_data[0]
    row[1].text = row_data[1]
    row[2].text = row_data[2]

doc.add_paragraph()
doc.add_paragraph('Admin Features (System Admin only):')
doc.add_paragraph('Check User Access:')
doc.add_paragraph('1. Select any user from the dropdown')
doc.add_paragraph('2. See their access level and accessible audit count')
doc.add_paragraph('3. Useful for troubleshooting access issues')
doc.add_paragraph('Admin Override:')
doc.add_paragraph('Use this for emergency access situations:')
doc.add_paragraph('1. Enter the Audit ID')
doc.add_paragraph('2. Select the user who needs access')
doc.add_paragraph('3. Enter a reason (required for audit trail)')
doc.add_paragraph('4. Click "Apply Admin Override"')
doc.add_paragraph('Warning: Admin overrides are logged and audited. Use only for legitimate emergency access.')
doc.add_paragraph('Department Access Summary:')
doc.add_paragraph('Shows how each role type accesses audits:')
doc.add_paragraph('- System Administrators → Full System Access')
doc.add_paragraph('- Audit Managers → Department + Assigned')
doc.add_paragraph('- Auditors → Assigned Audits Only')
doc.add_paragraph('- Department Staff → Department Audits')

# Sub-Feature 4: Role Matrix
doc.add_heading('Sub-Feature 4: Role Matrix', 3)
doc.add_paragraph('What it does: Define custom roles with specific permissions for fine-grained access control.')
doc.add_paragraph('Overview Cards show:')
doc.add_paragraph('- Total Roles in the system')
doc.add_paragraph('- System Roles (high privilege)')
doc.add_paragraph('- Global Roles (cross-department)')
doc.add_paragraph('- Audit Roles (audit-specific)')
doc.add_paragraph('Role Categories:')

table6 = doc.add_table(rows=5, cols=3)
table6.style = 'Table Grid'
hdr6 = table6.rows[0].cells
hdr6[0].text = 'Category'
hdr6[1].text = 'Purpose'
hdr6[2].text = 'Example'
cat_rows = [
    ['System', 'High-level system administration', 'System Administrator'],
    ['Audit', 'Audit-related activities', 'Lead Auditor, Auditor'],
    ['Business', 'Business operations', 'Department Manager'],
    ['Compliance', 'Compliance activities', 'Compliance Officer']
]
for i, row_data in enumerate(cat_rows):
    row = table6.rows[i+1].cells
    row[0].text = row_data[0]
    row[1].text = row_data[1]
    row[2].text = row_data[2]

doc.add_paragraph()
doc.add_paragraph('Role Scope:')
doc.add_paragraph('- Global Role - Applies across all departments')
doc.add_paragraph('- Department Role - Limited to specific department')
doc.add_paragraph('Creating a New Role:')
doc.add_paragraph('1. Click "Create Role"')
doc.add_paragraph('2. Enter role name and description')
doc.add_paragraph('3. Select category (System, Audit, Business, Compliance)')
doc.add_paragraph('4. Check "Global Role" if it applies across departments')
doc.add_paragraph('5. Select permissions (see below)')
doc.add_paragraph('6. Click save')

doc.add_paragraph('Available Permissions:')
doc.add_paragraph('Audit Management:')
doc.add_paragraph('- Create Audits')
doc.add_paragraph('- View All Audits')
doc.add_paragraph('- View Assigned Audits')
doc.add_paragraph('- Edit Audits')
doc.add_paragraph('- Delete Audits')
doc.add_paragraph('- Approve Reports')
doc.add_paragraph('System Management:')
doc.add_paragraph('- Manage Users')
doc.add_paragraph('- Manage Departments')
doc.add_paragraph('- View Analytics')
doc.add_paragraph('- Export Data')
doc.add_paragraph('Risk & CAPA:')
doc.add_paragraph('- Create Risks')
doc.add_paragraph('- Assess Risks')
doc.add_paragraph('- Approve Risk Treatments')
doc.add_paragraph('- Create CAPA')
doc.add_paragraph('- Assign CAPA')
doc.add_paragraph('- Close CAPA')
doc.add_paragraph('Document Control:')
doc.add_paragraph('- Upload Documents')
doc.add_paragraph('- Approve Documents')
doc.add_paragraph('- Archive Documents')
doc.add_paragraph('Asset & Vendor:')
doc.add_paragraph('- Manage Assets')
doc.add_paragraph('- Assign Assets')
doc.add_paragraph('- Manage Vendors')
doc.add_paragraph('- Evaluate Vendors')
doc.add_paragraph('Managing Existing Roles:')
doc.add_paragraph('- View role details (click eye icon)')
doc.add_paragraph('- Edit role permissions (click edit icon)')
doc.add_paragraph('- Delete role (click trash icon) - use with caution')
doc.add_paragraph('Role Table shows:')
doc.add_paragraph('- Role name and description')
doc.add_paragraph('- Category badge')
doc.add_paragraph('- Scope (Global/Department)')
doc.add_paragraph('- Number of permissions')
doc.add_paragraph('- Active/Inactive status')

# How Access Control Relates to Other Features
doc.add_heading('How Access Control Relates to Other Features', 3)
doc.add_paragraph('Connection to other features:')

table7 = doc.add_table(rows=9, cols=2)
table7.style = 'Table Grid'
hdr7 = table7.rows[0].cells
hdr7[0].text = 'Feature'
hdr7[1].text = 'How Access Control Affects It'
conn_rows = [
    ['Audits', 'Determines who can create, view, edit audits'],
    ['Team Assignment', 'Controls who can be assigned to audit teams'],
    ['Workflows', 'Determines who can approve workflow steps'],
    ['Reports', 'Controls who can generate and approve reports'],
    ['CAPA', 'Determines who can create, assign, close CAPAs'],
    ['Documents', 'Controls document upload and approval rights'],
    ['Analytics', 'Limits who can view analytics data'],
    ['Risk Assessment', 'Controls who can create and assess risks']
]
for i, row_data in enumerate(conn_rows):
    row = table7.rows[i+1].cells
    row[0].text = row_data[0]
    row[1].text = row_data[1]

doc.add_paragraph()

# Best Practices for Access Control
doc.add_heading('Best Practices for Access Control', 3)
doc.add_paragraph('1. Follow least privilege - Give users only the permissions they need')
doc.add_paragraph('2. Use temporary assignments - For short-term access needs, set expiry dates')
doc.add_paragraph('3. Document override reasons - Always explain why emergency access was granted')
doc.add_paragraph('4. Review roles regularly - Audit role assignments quarterly')
doc.add_paragraph('5. Verify team competency - Ensure audit team members are qualified')
doc.add_paragraph('6. Segregate duties - Don\'t let one person control entire processes')

# Gap Analysis
doc.add_heading('Gap Analysis', 1)
doc.add_paragraph('What it does: Compare your organization against ISO framework requirements.')
doc.add_paragraph('Features:')
doc.add_paragraph('- Select ISO framework (ISO 27001, ISO 9001, etc.)')
doc.add_paragraph('- Assess compliance for each clause')
doc.add_paragraph('- Identify gaps')
doc.add_paragraph('- Create remediation plans')
doc.add_paragraph('- Link gaps to CAPA items')
doc.add_paragraph('- Track closure progress')
doc.add_paragraph('Compliance Dashboard shows:')
doc.add_paragraph('- Overall compliance percentage')
doc.add_paragraph('- Gaps by severity')
doc.add_paragraph('- Remediation progress')
doc.add_paragraph('ISO Reference: Multiple frameworks supported')

# How to Conduct an Audit (Step-by-Step)
doc.add_heading('How to Conduct an Audit (Step-by-Step)', 1)

doc.add_heading('Step 1: Create the Audit', 2)
doc.add_paragraph('1. Go to Audits → Create Audit')
doc.add_paragraph('2. Enter audit title, year, scope')
doc.add_paragraph('3. Select department and risk rating')
doc.add_paragraph('4. Click Create')

doc.add_heading('Step 2: Initiate (ISO 19011 Clause 6.2)', 2)
doc.add_paragraph('1. Open the audit → Initiate tab')
doc.add_paragraph('2. Define audit objectives')
doc.add_paragraph('3. Set audit criteria (standards to audit against)')
doc.add_paragraph('4. Define detailed scope')
doc.add_paragraph('5. Select methodology')
doc.add_paragraph('6. Confirm feasibility')

doc.add_heading('Step 3: Assign Team', 2)
doc.add_paragraph('1. Go to Team tab')
doc.add_paragraph('2. Assign lead auditor')
doc.add_paragraph('3. Add team members with roles')
doc.add_paragraph('4. Verify competency')

doc.add_heading('Step 4: Prepare (ISO 19011 Clause 6.3)', 2)
doc.add_paragraph('1. Go to Prepare tab')
doc.add_paragraph('2. Create audit checklists')
doc.add_paragraph('3. Send document requests to auditees')
doc.add_paragraph('4. Conduct risk assessment')
doc.add_paragraph('5. Plan interview schedule')

doc.add_heading('Step 5: Execute (ISO 19011 Clause 6.4)', 2)
doc.add_paragraph('1. Click Start Execution')
doc.add_paragraph('2. Go to Evidence tab to upload evidence')
doc.add_paragraph('3. Use Work Program to track procedures')
doc.add_paragraph('4. Document observations')
doc.add_paragraph('5. Conduct interviews and record notes')
doc.add_paragraph('6. Go to Findings to document non-conformities')

doc.add_heading('Step 6: Report (ISO 19011 Clause 6.5)', 2)
doc.add_paragraph('1. Click Start Reporting')
doc.add_paragraph('2. Go to Report tab')
doc.add_paragraph('3. Generate report (AI-assisted)')
doc.add_paragraph('4. Review and edit content')
doc.add_paragraph('5. Submit for approval')
doc.add_paragraph('6. Create workflow for sign-off')

doc.add_heading('Step 7: Follow-up (ISO 19011 Clause 6.6)', 2)
doc.add_paragraph('1. Click Start Follow-up')
doc.add_paragraph('2. Go to Follow-up tab')
doc.add_paragraph('3. Create follow-up items for each finding')
doc.add_paragraph('4. Assign responsible persons')
doc.add_paragraph('5. Set due dates')
doc.add_paragraph('6. Track completion')
doc.add_paragraph('7. Verify effectiveness')

doc.add_heading('Step 8: Close Audit', 2)
doc.add_paragraph('1. Ensure all follow-ups are complete')
doc.add_paragraph('2. Click Close Audit')
doc.add_paragraph('3. Audit moves to Closed status')

# Tips for Success
doc.add_heading('Tips for Success', 1)
doc.add_paragraph('1. Always upload evidence - Every finding should have supporting evidence')
doc.add_paragraph('2. Use the workflow system - Get proper approvals for reports')
doc.add_paragraph('3. Track CAPA items - Don\'t let corrective actions slip')
doc.add_paragraph('4. Review the dashboard daily - Stay on top of overdue items')
doc.add_paragraph('5. Document everything - The system maintains full audit trail')

# How Features Relate to Each Other
doc.add_heading('How Features Relate to Each Other', 1)
doc.add_paragraph('Understanding how features connect helps you use the system more effectively.')

# Detailed Feature Connections
doc.add_heading('Detailed Feature Connections', 2)

doc.add_heading('Audits → Everything', 3)
doc.add_paragraph('Audits are the central hub. Almost everything connects to an audit:')

table8 = doc.add_table(rows=9, cols=2)
table8.style = 'Table Grid'
hdr8 = table8.rows[0].cells
hdr8[0].text = 'Feature'
hdr8[1].text = 'Connection to Audit'
feat_rows = [
    ['Evidence', 'Evidence is uploaded FOR a specific audit'],
    ['Findings', 'Findings are discovered DURING an audit'],
    ['Reports', 'Reports summarize AN audit'],
    ['Workflows', 'Workflows approve audit reports'],
    ['CAPA', 'CAPA items fix problems found in audits'],
    ['Follow-ups', 'Follow-ups track actions from audit findings'],
    ['Risk Assessment', 'Risks can be identified during audits'],
    ['Gap Analysis', 'Gaps are assessed as part of audits']
]
for i, row_data in enumerate(feat_rows):
    row = table8.rows[i+1].cells
    row[0].text = row_data[0]
    row[1].text = row_data[1]

doc.add_paragraph()

doc.add_heading('Findings → CAPA → Follow-ups', 3)
doc.add_paragraph('This is the corrective action chain:')
doc.add_paragraph('1. Finding discovered → A problem is found during audit execution')
doc.add_paragraph('2. CAPA created → A corrective/preventive action is defined to fix the finding')
doc.add_paragraph('3. Follow-up assigned → Tasks are created to track CAPA implementation')
doc.add_paragraph('4. Verification → Follow-up confirms the fix worked')
doc.add_paragraph('Example flow:')
doc.add_paragraph('Finding: "Password policy not enforced"')
doc.add_paragraph('CAPA: "Implement password complexity requirements"')
doc.add_paragraph('Follow-up: "IT to configure Active Directory by Jan 15"')
doc.add_paragraph('Verification: "Tested - passwords now require 12 characters"')

doc.add_heading('Risk Assessment → CAPA', 3)
doc.add_paragraph('Risks can trigger preventive actions:')
doc.add_paragraph('- High/Critical risks should have CAPA items to mitigate them')
doc.add_paragraph('- Risk controls link to CAPA for implementation tracking')
doc.add_paragraph('- When you create a risk, you can directly create a linked CAPA')

doc.add_heading('Risk Assessment → Assets', 3)
doc.add_paragraph('Risks are often tied to assets:')
doc.add_paragraph('- Each asset can have multiple risk assessments')
doc.add_paragraph('- Asset criticality affects risk scoring')
doc.add_paragraph('- Risk matrix shows risks by asset category')

doc.add_heading('Gap Analysis → CAPA', 3)
doc.add_paragraph('Compliance gaps need corrective actions:')
doc.add_paragraph('- Each gap identified can generate a CAPA item')
doc.add_paragraph('- Gap closure is tracked through CAPA completion')
doc.add_paragraph('- Compliance score improves as CAPAs close')

doc.add_heading('Documents → Audits', 3)
doc.add_paragraph('Documents support the audit process:')
doc.add_paragraph('- Policies and procedures are audit criteria')
doc.add_paragraph('- Evidence documents are uploaded during execution')
doc.add_paragraph('- Audit reports are stored as controlled documents')

doc.add_heading('Workflows → Reports', 3)
doc.add_paragraph('Workflows manage report approval:')
doc.add_paragraph('- When a report is ready, create a workflow')
doc.add_paragraph('- Workflow routes report through approvers')
doc.add_paragraph('- Each department head reviews and signs')
doc.add_paragraph('- Report is published only after workflow completes')

doc.add_heading('Vendors → Risk Assessment', 3)
doc.add_paragraph('Third-party risks:')
doc.add_paragraph('- Vendors have risk ratings')
doc.add_paragraph('- High-risk vendors should have risk assessments')
doc.add_paragraph('- Vendor audits can be scheduled based on risk')

doc.add_heading('Assets → Audits', 3)
doc.add_paragraph('Assets define audit scope:')
doc.add_paragraph('- Select which assets are in scope for an audit')
doc.add_paragraph('- Asset owners are notified of audits')
doc.add_paragraph('- Findings can be linked to specific assets')

doc.add_heading('Departments → Users → Audits', 3)
doc.add_paragraph('Organizational structure:')
doc.add_paragraph('- Users belong to departments')
doc.add_paragraph('- Departments are assigned to audits (as auditee)')
doc.add_paragraph('- Department heads approve findings for their area')
doc.add_paragraph('- Workflow steps route to departments')

doc.add_heading('Dashboard → Everything', 3)
doc.add_paragraph('Dashboard aggregates all data:')
doc.add_paragraph('- Shows audit counts from Audits')
doc.add_paragraph('- Shows finding counts from Findings')
doc.add_paragraph('- Shows risk heatmap from Risk Assessment')
doc.add_paragraph('- Shows compliance scores from Gap Analysis')
doc.add_paragraph('- Shows CAPA status from CAPA Management')
doc.add_paragraph('- Shows overdue items from Follow-ups')

doc.add_heading('Analytics → Everything', 3)
doc.add_paragraph('Analytics provides trends across:')
doc.add_paragraph('- Audit completion rates over time')
doc.add_paragraph('- Finding trends by severity')
doc.add_paragraph('- CAPA closure rates')
doc.add_paragraph('- Compliance improvement')
doc.add_paragraph('- Risk reduction')

# Common Workflows Between Features
doc.add_heading('Common Workflows Between Features', 2)

doc.add_paragraph('Workflow 1: Complete Audit Cycle')
doc.add_paragraph('Planning → Audit Created → Team Assigned → Preparation → Execution (Evidence + Findings) → Report → Workflow Approval → CAPA Created → Follow-ups → Verification → Audit Closed')

doc.add_paragraph('Workflow 2: Risk-Based Audit')
doc.add_paragraph('Risk Assessment → High Risk Identified → Audit Planned → Audit Executed → Controls Tested → CAPA for Weak Controls → Risk Re-assessed → Risk Reduced')

doc.add_paragraph('Workflow 3: Compliance Gap Closure')
doc.add_paragraph('Gap Analysis → Gap Identified → CAPA Created → Actions Implemented → Evidence Uploaded → Gap Verified → Compliance Score Updated')

doc.add_paragraph('Workflow 4: Document Control')
doc.add_paragraph('Document Uploaded (Draft) → Review Workflow Created → Reviewers Approve → Document Active → Linked to Audit as Criteria → Expiry Tracked')

# Quick Reference: What Links to What
doc.add_heading('Quick Reference: What Links to What', 2)

table9 = doc.add_table(rows=9, cols=2)
table9.style = 'Table Grid'
hdr9 = table9.rows[0].cells
hdr9[0].text = 'If you\'re in...'
hdr9[1].text = 'You can link to...'
link_rows = [
    ['Audit', 'Department, Users (team), Assets, Documents'],
    ['Finding', 'Audit, CAPA, Follow-up, Evidence'],
    ['CAPA', 'Audit, Finding, Risk, Gap, User (assignee)'],
    ['Risk', 'Audit, Asset, CAPA, Controls'],
    ['Gap', 'Framework, Audit, CAPA, Department'],
    ['Document', 'Department, Audit (as evidence)'],
    ['Workflow', 'Audit, Report, Users (approvers)'],
    ['Follow-up', 'Audit, Finding, User (assignee)']
]
for i, row_data in enumerate(link_rows):
    row = table9.rows[i+1].cells
    row[0].text = row_data[0]
    row[1].text = row_data[1]

doc.add_paragraph()

# Data Flow Summary
doc.add_heading('Data Flow Summary', 2)
doc.add_paragraph('Input sources:')
doc.add_paragraph('- Users create audits, risks, documents')
doc.add_paragraph('- Auditors upload evidence and findings')
doc.add_paragraph('- System generates reports')
doc.add_paragraph('Processing:')
doc.add_paragraph('- Workflows route approvals')
doc.add_paragraph('- CAPA tracks corrective actions')
doc.add_paragraph('- Follow-ups monitor completion')
doc.add_paragraph('Output/Reporting:')
doc.add_paragraph('- Dashboard shows real-time status')
doc.add_paragraph('- Analytics shows trends')
doc.add_paragraph('- Reports document results')

# Need Help?
doc.add_heading('Need Help?', 1)
doc.add_paragraph('Contact your System Administrator for:')
doc.add_paragraph('- Password resets')
doc.add_paragraph('- Role changes')
doc.add_paragraph('- Access issues')
doc.add_paragraph('- Training requests')

doc.add_paragraph('')
doc.add_paragraph('This system follows ISO 19011:2018 Guidelines for auditing management systems')

# Save the document
doc.save('Galaxy_Audit_System_Guide.docx')
print('Document created successfully: Galaxy_Audit_System_Guide.docx')
