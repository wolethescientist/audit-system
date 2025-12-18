"""
Simple Report Generation Service
"""

import os
import io
import logging
import base64
from typing import Dict, Any
from datetime import datetime
from uuid import uuid4
from sqlalchemy.orm import Session
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from docx import Document
import markdown

from app.models import Audit, AuditReport, ReportStatus
from app.services.gemini_service import GeminiAIService

logger = logging.getLogger(__name__)


class ReportGenerationService:
    """Simple service for generating audit reports."""
    
    def __init__(self):
        self.gemini_service = None  # Lazy init
    
    def _get_gemini_service(self):
        if self.gemini_service is None:
            self.gemini_service = GeminiAIService()
        return self.gemini_service
    
    async def generate_report(self, audit_id: str, db: Session, user_id: str) -> Dict[str, Any]:
        """Generate audit report."""
        
        # Validate audit exists
        audit = db.query(Audit).filter(Audit.id == audit_id).first()
        if not audit:
            raise ValueError(f"Audit with ID {audit_id} not found")
        
        logger.info(f"Generating report for audit: {audit.title}")
        
        # Generate report content using Gemini
        gemini = self._get_gemini_service()
        ai_result = await gemini.generate_audit_report(audit_id, db)
        
        # Save report to database
        report_record = self._save_report(
            audit_id=audit_id,
            content=ai_result["content"],
            user_id=user_id,
            db=db
        )
        
        # Generate download formats
        html_content = self._markdown_to_html(ai_result["content"], audit)
        pdf_content = self._generate_pdf(ai_result["content"], audit)
        docx_content = self._generate_docx(ai_result["content"], audit)
        
        return {
            "report_id": str(report_record.id),
            "audit_id": audit_id,
            "status": "completed",
            "content": ai_result["content"],
            "html_content": html_content,
            "generation_date": ai_result["generation_date"],
            "iso_compliance_validated": ai_result["iso_compliance_validated"],
            "validation_notes": ai_result["validation_notes"],
            "word_count": ai_result["word_count"],
            "sections_generated": ai_result["sections_generated"],
            "download_files": {
                "pdf_content": pdf_content,
                "docx_content": docx_content
            },
            "supported_formats": ["pdf", "docx"]
        }
    
    def _save_report(self, audit_id: str, content: str, user_id: str, db: Session) -> AuditReport:
        """Save or update report in database."""
        
        existing = db.query(AuditReport).filter(AuditReport.audit_id == audit_id).first()
        
        if existing:
            existing.content = content
            existing.version += 1
            existing.created_by_id = user_id
            existing.created_at = datetime.utcnow()
            db.commit()
            db.refresh(existing)
            return existing
        
        report = AuditReport(
            id=uuid4(),
            audit_id=audit_id,
            version=1,
            content=content,
            status=ReportStatus.DRAFT,
            created_by_id=user_id
        )
        db.add(report)
        db.commit()
        db.refresh(report)
        return report
    
    def _markdown_to_html(self, content: str, audit: Audit) -> str:
        """Convert markdown to styled HTML."""
        
        html_body = markdown.markdown(content, extensions=['tables'])
        
        return f"""<!DOCTYPE html>
<html>
<head>
    <title>Audit Report - {audit.title}</title>
    <style>
        body {{ font-family: Arial, sans-serif; max-width: 900px; margin: 0 auto; padding: 20px; }}
        h1 {{ color: #2c3e50; border-bottom: 2px solid #3498db; padding-bottom: 10px; }}
        h2 {{ color: #34495e; }}
        table {{ border-collapse: collapse; width: 100%; margin: 20px 0; }}
        th, td {{ border: 1px solid #ddd; padding: 10px; text-align: left; }}
        th {{ background: #f5f5f5; }}
    </style>
</head>
<body>
{html_body}
</body>
</html>"""
    
    def _generate_pdf(self, content: str, audit: Audit) -> str:
        """Generate PDF and return as base64."""
        
        try:
            buffer = io.BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=A4)
            styles = getSampleStyleSheet()
            story = []
            
            # Add content
            for line in content.split('\n'):
                line = line.strip()
                if not line:
                    story.append(Spacer(1, 12))
                elif line.startswith('# '):
                    story.append(Paragraph(line[2:], styles['Heading1']))
                elif line.startswith('## '):
                    story.append(Paragraph(line[3:], styles['Heading2']))
                elif line.startswith('### '):
                    story.append(Paragraph(line[4:], styles['Heading3']))
                elif line.startswith('- '):
                    story.append(Paragraph(f"• {line[2:]}", styles['Normal']))
                else:
                    # Clean markdown formatting
                    clean_line = line.replace('**', '').replace('*', '').replace('|', ' ')
                    if clean_line.strip():
                        story.append(Paragraph(clean_line, styles['Normal']))
            
            doc.build(story)
            buffer.seek(0)
            return base64.b64encode(buffer.getvalue()).decode('utf-8')
            
        except Exception as e:
            logger.error(f"PDF generation error: {e}")
            return ""
    
    def _generate_docx(self, content: str, audit: Audit) -> str:
        """Generate Word doc and return as base64."""
        
        try:
            doc = Document()
            
            for line in content.split('\n'):
                line = line.strip()
                if not line:
                    doc.add_paragraph("")
                elif line.startswith('# '):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith('## '):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith('### '):
                    doc.add_heading(line[4:], level=3)
                elif line.startswith('- '):
                    doc.add_paragraph(line[2:], style='List Bullet')
                else:
                    # Clean markdown
                    clean_line = line.replace('**', '').replace('*', '').replace('|', ' ')
                    if clean_line.strip() and not clean_line.startswith('---'):
                        doc.add_paragraph(clean_line)
            
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return base64.b64encode(buffer.getvalue()).decode('utf-8')
            
        except Exception as e:
            logger.error(f"DOCX generation error: {e}")
            return ""

    async def _generate_pdf_content(self, content: str, audit: Audit) -> str:
        """Wrapper for PDF generation."""
        return self._generate_pdf(content, audit)
    
    async def _generate_docx_content(self, content: str, audit: Audit) -> str:
        """Wrapper for DOCX generation."""
        return self._generate_docx(content, audit)
    
    async def _generate_csv_content(self, audit: Audit, report_id: str) -> str:
        """Generate simple CSV."""
        return f"Field,Value\nAudit ID,{audit.id}\nTitle,{audit.title}\nYear,{audit.year}\nStatus,{audit.status.value}"
    
    async def _generate_html_content(self, content: str, audit: Audit) -> str:
        """Wrapper for HTML generation."""
        return self._markdown_to_html(content, audit)
    
    async def _generate_markdown_content(self, content: str, audit: Audit) -> str:
        """Return markdown with metadata."""
        return f"---\ntitle: {audit.title}\nyear: {audit.year}\n---\n\n{content}"
